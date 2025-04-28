# backend/document_generator.py

import io
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_docx_document(question, answer, sources=None):
    """
    Creates a DOCX document with the question, answer, and sources.
    
    Args:
        question: The DDQ question
        answer: The AI-generated answer
        sources: List of source documents used (optional)
        
    Returns:
        A BytesIO object containing the DOCX document
    """
    # Create a new Document
    doc = Document()
    
    # Set up document styles
    styles = doc.styles
    
    # Modify the Normal style
    style_normal = styles['Normal']
    font = style_normal.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Create a style for the title
    style_title = styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)
    style_title.base_style = styles['Normal']
    style_title.font.name = 'Calibri'
    style_title.font.size = Pt(16)
    style_title.font.bold = True
    style_title.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue
    
    # Create a style for headings
    style_heading = styles.add_style('Heading', WD_STYLE_TYPE.PARAGRAPH)
    style_heading.base_style = styles['Normal']
    style_heading.font.name = 'Calibri'
    style_heading.font.size = Pt(14)
    style_heading.font.bold = True
    
    # Create a style for sources
    style_source = styles.add_style('Source', WD_STYLE_TYPE.PARAGRAPH)
    style_source.base_style = styles['Normal']
    style_source.font.name = 'Calibri'
    style_source.font.size = Pt(10)
    style_source.font.italic = True
    
    # Add a title
    title = doc.add_paragraph("DDQ Response", style='Title')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date and time
    date_time = doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    date_time.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Add a horizontal line
    doc.add_paragraph("_" * 80)
    
    # Add the question
    doc.add_paragraph("Question:", style='Heading')
    doc.add_paragraph(question)
    
    # Add some space
    doc.add_paragraph()
    
    # Add the answer
    doc.add_paragraph("Answer:", style='Heading')
    doc.add_paragraph(answer)
    
    # Add sources if provided
    if sources and len(sources) > 0:
        # Add some space
        doc.add_paragraph()
        
        # Add sources heading
        doc.add_paragraph("Sources:", style='Heading')
        
        # Add each source
        for source in sources:
            doc.add_paragraph(f"• {source}", style='Source')
    
    # Add a footer
    section = doc.sections[0]
    footer = section.footer
    footer_text = footer.paragraphs[0]
    footer_text.text = "Hudson Advisors DDQ Assistant"
    footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save the document to a BytesIO object
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    return docx_bytes

def generate_and_upload_docx(question, answer, sources, blob_service):
    """
    Generates a DOCX document and uploads it to Azure Blob Storage.
    
    Args:
        question: The DDQ question
        answer: The AI-generated answer
        sources: List of source documents used
        blob_service: Instance of BlobStorageService
        
    Returns:
        The URL of the uploaded document if successful, None otherwise
    """
    try:
        # Generate a unique blob name
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        # Create a safe filename from the question (first 30 chars)
        safe_question = "".join(c if c.isalnum() else "_" for c in question[:30])
        blob_name = f"ddq_responses/{timestamp}_{safe_question}.docx"
        
        # Create the DOCX document
        docx_bytes = create_docx_document(question, answer, sources)
        
        # Upload the document to Blob Storage
        document_url = blob_service.upload_document(
            docx_bytes.getvalue(), 
            blob_name,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        return document_url
    except Exception as e:
        print(f"Error generating and uploading document: {e}")
        return None

# Example usage (for testing purposes)
if __name__ == "__main__":
    try:
        from blob_storage_service import BlobStorageService
        
        # This requires environment variables to be set
        blob_service = BlobStorageService()
        
        # Example question and answer
        question = "What is the fund's ESG policy?"
        answer = "The fund's ESG policy emphasizes responsible investment practices across all portfolio companies. It includes regular ESG assessments, carbon footprint monitoring, and adherence to international standards such as the UN Principles for Responsible Investment."
        sources = ["ESG Policy.pdf", "Q&A Bank.xlsx"]
        
        # Generate and upload the document
        document_url = generate_and_upload_docx(question, answer, sources, blob_service)
        
        if document_url:
            print(f"Document uploaded successfully: {document_url}")
        else:
            print("Failed to upload document.")
    except Exception as e:
        print(f"Error in example: {e}")
        print("Please ensure Azure Blob Storage environment variables are set.")
